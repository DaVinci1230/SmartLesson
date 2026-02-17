"""
TOS File Parser Service

Parses TOS (Table of Specifications) from various file formats:
- JSON: Direct structured format
- PDF: Extracts TOS matrix and learning outcomes
- DOCX: Extracts TOS data from tables

This service bridges external TOS uploads with the internal TQS generation pipeline.

Key Output Format (Normalized):
{
    "learning_outcomes": [
        {
            "id": int,
            "text": str,
            "hours": float (optional)
        }
    ],
    "bloom_distribution": {
        "Remember": int,      # percentage or count
        "Understand": int,
        "Apply": int,
        "Analyze": int,
        "Evaluate": int,
        "Create": int
    },
    "tos_matrix": {
        "Remember": {
            "outcome_0": int,  # number of items
            "outcome_1": int,
            ...
        },
        "Understand": {...},
        ...
    },
    "total_items": int,
    "total_points": int (optional),
    "metadata": {
        "course_code": str (optional),
        "course_title": str (optional),
        "semester": str (optional),
        "file_name": str,
        "parsed_at": str
    }
}
"""

import json
import logging
from typing import Dict, List, Any, Optional, Tuple
import io
from datetime import datetime
import pandas as pd

# Try importing optional dependencies
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False
    PyPDF2 = None

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False
    openpyxl = None

logger = logging.getLogger(__name__)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

BLOOM_LEVELS = [
    "Remember",
    "Understand",
    "Apply",
    "Analyze",
    "Evaluate",
    "Create"
]


class TOSParsingError(Exception):
    """Custom exception for TOS parsing errors"""
    pass


class TOSFileParser:
    """
    Main parser class for converting various TOS formats to normalized structure.
    """

    def __init__(self):
        self.supported_formats = [".json"]
        if HAS_PYPDF2:
            self.supported_formats.append(".pdf")
        if HAS_DOCX:
            self.supported_formats.append(".docx")
        if HAS_OPENPYXL:
            self.supported_formats.append(".xlsx")

    def parse(
        self,
        file_content: bytes,
        file_name: str,
        file_type: str = None
    ) -> Dict[str, Any]:
        """
        Parse a TOS file and return normalized structure.

        Args:
            file_content: Raw bytes of the file
            file_name: Original filename
            file_type: File type (json, pdf, docx). Auto-detected if not provided.

        Returns:
            Normalized TOS structure

        Raises:
            TOSParsingError: If file cannot be parsed
        """
        if file_type is None:
            file_type = self._detect_file_type(file_name)

        if file_type == "json":
            return self._parse_json(file_content, file_name)
        elif file_type == "pdf":
            if not HAS_PYPDF2:
                raise TOSParsingError(
                    "PDF parsing not available. Install PyPDF2: pip install PyPDF2"
                )
            return self._parse_pdf(file_content, file_name)
        elif file_type == "docx":
            if not HAS_DOCX:
                raise TOSParsingError(
                    "DOCX parsing not available. Install python-docx: pip install python-docx"
                )
            return self._parse_docx(file_content, file_name)
        elif file_type == "xlsx":
            if not HAS_OPENPYXL:
                raise TOSParsingError(
                    "Excel parsing not available. Install openpyxl: pip install openpyxl"
                )
            return self._parse_xlsx(file_content, file_name)
        else:
            raise TOSParsingError(
                f"Unsupported file type: {file_type}. "
                f"Supported: {', '.join(self.supported_formats)}"
            )

    def _detect_file_type(self, file_name: str) -> str:
        """Detect file type from filename extension"""
        if not file_name:
            raise TOSParsingError("File name cannot be empty")

        file_name_lower = file_name.lower()
        if file_name_lower.endswith(".json"):
            return "json"
        elif file_name_lower.endswith(".pdf"):
            return "pdf"
        elif file_name_lower.endswith((".docx", ".doc")):
            return "docx"
        elif file_name_lower.endswith(".xlsx"):
            return "xlsx"
        else:
            raise TOSParsingError(f"Cannot detect file type from: {file_name}")

    def _parse_json(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """
        Parse JSON TOS file.

        Expected JSON structure:
        {
            "learning_outcomes": [...],
            "bloom_distribution": {...},
            "tos_matrix": {...},
            "total_items": int,
            "metadata": {...}
        }
        """
        try:
            content_str = file_content.decode("utf-8")
            data = json.loads(content_str)
        except (UnicodeDecodeError, json.JSONDecodeError) as e:
            raise TOSParsingError(f"Invalid JSON file: {str(e)}")

        # Validate and normalize the JSON structure
        return self._validate_and_normalize_tos(data, file_name)

    def _parse_pdf(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """
        Parse PDF TOS file.

        Expected: PDF containing a well-structured table with:
        - Column headers: Learning Outcomes + Bloom levels
        - Rows: One per outcome with item counts per Bloom level
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Extract text from all pages
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"

            if not text.strip():
                raise TOSParsingError("PDF contains no extractable text")

            # Try to parse tables from PDF text
            tos_data = self._extract_tos_from_pdf_text(text)
            tos_data["metadata"]["file_name"] = file_name

            return self._validate_and_normalize_tos(tos_data, file_name)

        except Exception as e:
            raise TOSParsingError(f"PDF parsing failed: {str(e)}")

    def _extract_tos_from_pdf_text(self, text: str) -> Dict[str, Any]:
        """
        Extract TOS structure from PDF text.

        This is a basic implementation. For production use, consider
        using pypdf or pdfplumber for better table extraction.
        """
        logger.warning(
            "PDF parsing is complex. For best results, use direct JSON format. "
            "Basic text extraction may miss table structure."
        )

        # Initialize basic structure
        tos_data = {
            "learning_outcomes": [],
            "bloom_distribution": {},
            "tos_matrix": {},
            "total_items": 0,
            "metadata": {
                "parsing_method": "text_extraction_basic"
            }
        }

        # This is a placeholder. In production, you would:
        # 1. Use pdfplumber for accurate table extraction
        # 2. Or ask user for JSON format which is more reliable
        logger.info("PDF text extraction completed. Please verify accuracy.")

        return tos_data

    def _parse_docx(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """
        Parse DOCX TOS file.

        Expected: DOCX containing a table with:
        - Header row: Learning Outcomes + Bloom levels + totals
        - Data rows: One per outcome with item counts
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            if not doc.tables:
                raise TOSParsingError("DOCX contains no tables")

            # Extract TOS from first table (main table often first)
            tos_data = self._extract_tos_from_docx_tables(doc.tables)
            tos_data["metadata"]["file_name"] = file_name

            return self._validate_and_normalize_tos(tos_data, file_name)

        except Exception as e:
            raise TOSParsingError(f"DOCX parsing failed: {str(e)}")

    def _extract_tos_from_docx_tables(self, tables: List) -> Dict[str, Any]:
        """
        Extract TOS from DOCX tables.

        Looks for table structure like:
        | Learning Outcome | Remember | Understand | Apply | ... | Total |
        | Outcome A        |    3     |     2      |   2   | ... |   10  |
        | Outcome B        |    2     |     3      |   3   | ... |   10  |
        """
        if not tables:
            raise TOSParsingError("No tables found in DOCX")

        # Use first table with most columns (usually the main TOS table)
        main_table = max(tables, key=lambda t: len(t.rows[0].cells))

        rows = main_table.rows
        if len(rows) < 2:
            raise TOSParsingError("Table must have header row and at least one data row")

        # Extract header
        header_cells = [cell.text.strip() for cell in rows[0].cells]
        header_lower = [h.lower() for h in header_cells]

        # Find bloom column indices
        bloom_columns = {}
        for bloom in BLOOM_LEVELS:
            for i, h in enumerate(header_lower):
                if bloom.lower() in h:
                    bloom_columns[bloom] = i
                    break

        if not bloom_columns:
            raise TOSParsingError(
                f"Could not find Bloom levels in table headers. Found: {header_cells}"
            )

        # Parse data rows
        outcomes = []
        tos_matrix = {bloom: {} for bloom in BLOOM_LEVELS}
        total_items = 0

        for row_idx in range(1, len(rows)):
            cells = [cell.text.strip() for cell in rows[row_idx].cells]

            if not cells[0]:  # Skip empty rows
                continue

            outcome_text = cells[0]
            outcome_id = row_idx - 1

            outcomes.append({
                "id": outcome_id,
                "text": outcome_text,
                "hours": 1.0  # Default, should be extracted from file if available
            })

            # Extract bloom counts for this outcome
            for bloom, col_idx in bloom_columns.items():
                try:
                    count = int(cells[col_idx]) if col_idx < len(cells) else 0
                    tos_matrix[bloom][outcome_id] = count
                    total_items += count
                except (ValueError, IndexError):
                    tos_matrix[bloom][outcome_id] = 0

        # Compute bloom distribution percentages
        bloom_distribution = {}
        for bloom in BLOOM_LEVELS:
            total_for_bloom = sum(tos_matrix[bloom].values())
            bloom_distribution[bloom] = total_for_bloom if total_items == 0 else int(
                (total_for_bloom / total_items) * 100
            )

        return {
            "learning_outcomes": outcomes,
            "bloom_distribution": bloom_distribution,
            "tos_matrix": tos_matrix,
            "total_items": total_items,
            "metadata": {
                "parsing_method": "docx_table_extraction"
            }
        }

    def _parse_xlsx(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """
        Parse Excel (.xlsx) TOS file.

        Expected: Excel sheet with:
        - Header row: Learning Outcomes + Bloom levels + totals
        - Data rows: One per outcome with item counts

        The parser looks for:
        - Column A: Learning Outcome text
        - Columns B-G: Remember, Understand, Apply, Analyze, Evaluate, Create
        - Column H (optional): Total items per outcome
        """
        try:
            excel_file = io.BytesIO(file_content)
            
            # Use pandas to detect the correct sheet and work with the data
            # This will also handle merged cells better
            tos_data = self._extract_tos_from_xlsx_file(excel_file, file_name)
            
            return self._validate_and_normalize_tos(tos_data, file_name)
            
        except Exception as e:
            raise TOSParsingError(f"Excel parsing failed: {str(e)}")

    def _extract_tos_from_xlsx_file(self, file_obj: io.BytesIO, file_name: str) -> Dict[str, Any]:
        """
        Extract TOS from Excel file using pandas for robust handling of merged cells
        and complex formatting.
        
        Args:
            file_obj: BytesIO object containing Excel file
            file_name: Original file name
            
        Returns:
            Normalized TOS structure
        """
        # Read all sheet names first
        try:
            sheet_names = pd.ExcelFile(file_obj).sheet_names
        except Exception:
            raise TOSParsingError("Could not read Excel file sheets")
        
        if not sheet_names:
            raise TOSParsingError("Excel file contains no sheets")
        
        # Use first sheet
        sheet_name = sheet_names[0]
        file_obj.seek(0)  # Reset file pointer
        
        # First pass: detect header row
        df_raw = pd.read_excel(file_obj, sheet_name=sheet_name, header=None)
        header_row_idx = self._detect_header_row_in_df(df_raw)
        
        if header_row_idx is None:
            raise TOSParsingError(
                f"Could not find header row with Bloom levels in Excel sheet. "
                f"Expected columns like: Learning Outcome, Remember, Understand, Apply, Analyze, Evaluate, Create"
            )
        
        # Second pass: read Excel with detected header
        file_obj.seek(0)
        df = pd.read_excel(
            file_obj,
            sheet_name=sheet_name,
            header=header_row_idx,
            dtype=str
        )
        
        # Extract TOS using the detected header
        tos_data = self._extract_tos_from_dataframe(df, header_row_idx)
        tos_data["metadata"]["file_name"] = file_name
        
        return tos_data
    
    def _detect_header_row_in_df(self, df: pd.DataFrame) -> Optional[int]:
        """
        Detect which row contains the header by looking for Bloom keywords.
        
        Args:
            df: DataFrame read with header=None
            
        Returns:
            Row index of header, or None if not found
        """
        for idx, row in df.iterrows():
            row_str = " ".join(str(cell).lower() for cell in row if pd.notna(cell))
            bloom_keywords = {'remember', 'understand', 'apply', 'analyze', 'evaluate', 'create'}
            found_blooms = sum(1 for keyword in bloom_keywords if keyword in row_str)
            if found_blooms >= 2:
                return idx
        return None

    def _extract_tos_from_dataframe(self, df: pd.DataFrame, header_row_idx: int) -> Dict[str, Any]:
        """
        Extract TOS structure from pandas DataFrame (after header is properly set).
        
        Handles:
        - Merged cells (via forward fill)
        - Blank separator rows
        - Formatted Excel sheets with totals/percentages
        - Indentation and whitespace
        """
        # Normalize column names (lowercase, strip whitespace)
        df.columns = [col.lower().strip() for col in df.columns]
        
        # Find outcome column (first column with 'outcome' or 'learning' in name)
        outcome_col = None
        for col in df.columns:
            if 'outcome' in col or 'learning' in col:
                outcome_col = col
                break
        if not outcome_col:
            # Use first column if no explicit outcome column found
            outcome_col = df.columns[0]
        
        # Handle merged cells - forward fill the outcome column
        df[outcome_col] = df[outcome_col].ffill()
        
        # Find Bloom columns
        bloom_columns = {}
        for bloom in BLOOM_LEVELS:
            for col in df.columns:
                if bloom.lower() in col.lower():
                    bloom_columns[bloom] = col
                    break
        
        if not bloom_columns:
            raise TOSParsingError(
                f"Could not find Bloom level columns. Found columns: {list(df.columns)}"
            )
        
        # Extract learning outcomes with robust validation
        outcomes = []
        tos_matrix = {bloom: {} for bloom in BLOOM_LEVELS}
        total_items = 0
        
        for idx, row in df.iterrows():
            outcome_text = str(row[outcome_col]).strip() if pd.notna(row[outcome_col]) else ""
            
            # Skip empty values
            if not outcome_text or outcome_text.lower() == 'nan':
                continue
            
            # Skip rows that are headers, totals, or purely numeric
            outcome_lower = outcome_text.lower()
            skip_keywords = ['total', 'percentage', 'percent', 'sum', 'subtotal', 'mean', 'average', 'grand total', outcome_col.lower()]
            if any(skip_word in outcome_lower for skip_word in skip_keywords):
                continue
            
            # Also skip if outcome is exactly a common non-outcome label
            if outcome_text.strip() in ['TOTAL', 'Percentage', 'Mean', 'Average', 'Grand Total', '%']:
                continue
            
            # Skip purely numeric rows (likely table artifacts)
            try:
                float(outcome_text)
                continue  # It's numeric, skip it
            except ValueError:
                pass  # It's text, keep it
            
            # Extract Bloom counts for this outcome
            outcome_id = len(outcomes)
            has_numeric_data = False
            
            for bloom, col in bloom_columns.items():
                try:
                    val = row[col] if col in row.index else ""
                    if pd.isna(val) or val == "" or str(val).lower() == 'nan':
                        count = 0
                    else:
                        count = int(float(str(val).strip()))
                    tos_matrix[bloom][outcome_id] = count
                    total_items += count
                    if count > 0:
                        has_numeric_data = True
                except (ValueError, TypeError):
                    tos_matrix[bloom][outcome_id] = 0
            
            # Only add outcomes that have at least some numeric data in Bloom columns
            if has_numeric_data:
                outcomes.append({
                    "id": outcome_id,
                    "text": outcome_text,
                    "hours": 1.0
                })
        
        if not outcomes:
            raise TOSParsingError(
                f"No valid learning outcomes found in Excel sheet. "
                f"Checked {len(df)} rows for textual outcomes with Bloom distribution data."
            )
        
        # Compute bloom distribution percentages
        bloom_distribution = {}
        for bloom in BLOOM_LEVELS:
            total_for_bloom = sum(tos_matrix[bloom].values())
            bloom_distribution[bloom] = total_for_bloom if total_items == 0 else int(
                (total_for_bloom / total_items) * 100
            )
        
        return {
            "learning_outcomes": outcomes,
            "bloom_distribution": bloom_distribution,
            "tos_matrix": tos_matrix,
            "total_items": total_items,
            "metadata": {
                "parsing_method": "xlsx_pandas_extraction"
            }
        }

    def _extract_tos_from_xlsx_sheet(self, sheet) -> Dict[str, Any]:
        """
        Extract TOS from Excel sheet with robust handling of complex formatting.

        Handles:
        - Title rows before actual header
        - Merged cells (using forward fill)
        - Blank separator rows
        - Formatted Excel sheets with totals/percentages
        - Indentation and whitespace

        Expected structure (ideally):
        | Learning Outcome | Remember | Understand | Apply | Analyze | Evaluate | Create |
        | Outcome A        |    3     |     2      |   2   |    1    |     1    |    1   |
        | Outcome B        |    2     |     3      |   3   |    2    |     2    |    1   |
        """
        # STEP 1: Read all raw cell values to detect header row
        rows = []
        for row in sheet.iter_rows(values_only=True):
            rows.append([str(cell).strip() if cell is not None else "" for cell in row])
        
        if len(rows) < 2:
            raise TOSParsingError("Sheet must have header row and at least one data row")
        
        # STEP 2: Detect header row dynamically by searching for Bloom keywords
        header_row_idx = None
        for row_idx, row in enumerate(rows):
            if not row:
                continue
            row_lower = [str(cell).lower() for cell in row]
            # Check if this row contains Bloom level keywords
            bloom_keywords = {'remember', 'understand', 'apply', 'analyze', 'evaluate', 'create'}
            found_blooms = sum(1 for cell in row_lower if any(keyword in cell for keyword in bloom_keywords))
            if found_blooms >= 2:  # At least 2 Bloom keywords to be confident
                header_row_idx = row_idx
                break
        
        if header_row_idx is None:
            raise TOSParsingError(
                f"Could not find header row with Bloom levels in Excel sheet. "
                f"Expected columns like: Learning Outcome, Remember, Understand, Apply, Analyze, Evaluate, Create"
            )
        
        # STEP 3: Read Excel with pandas starting from detected header row
        try:
            df = pd.read_excel(
                io.BytesIO(sheet.parent.filename or pd.io.common._stringify_path(sheet.parent)),
                sheet_name=sheet.title,
                header=header_row_idx,
                dtype=str  # Read all as strings for flexible processing
            )
        except Exception:
            # Fallback: manually construct dataframe from raw rows
            header = rows[header_row_idx]
            data_rows = rows[header_row_idx + 1:]
            df = pd.DataFrame(data_rows, columns=header)
        
        # STEP 4: Normalize column names (lowercase, strip whitespace)
        df.columns = [col.lower().strip() for col in df.columns]
        
        # STEP 5: Find first column (outcome column) - usually "learning outcome" or similar
        outcome_col = None
        for col in df.columns:
            if 'outcome' in col or 'learning' in col:
                outcome_col = col
                break
        if not outcome_col:
            # Use first column if no explicit outcome column found
            outcome_col = df.columns[0]
        
        # STEP 6: Handle merged cells - forward fill the outcome column
        df[outcome_col] = df[outcome_col].fillna(method='ffill')
        
        # STEP 7: Find Bloom columns
        bloom_columns = {}
        for bloom in BLOOM_LEVELS:
            for col in df.columns:
                if bloom.lower() in col.lower():
                    bloom_columns[bloom] = col
                    break
        
        if not bloom_columns:
            raise TOSParsingError(
                f"Could not find Bloom level columns in Excel. Found columns: {list(df.columns)}"
            )
        
        # STEP 8: Extract learning outcomes with validation
        outcomes = []
        tos_matrix = {bloom: {} for bloom in BLOOM_LEVELS}
        total_items = 0
        
        for idx, row in df.iterrows():
            outcome_text = str(row[outcome_col]).strip() if pd.notna(row[outcome_col]) else ""
            
            # Skip empty values
            if not outcome_text:
                continue
            
            # Skip rows that are headers, totals, or purely numeric
            outcome_lower = outcome_text.lower()
            if any(skip_word in outcome_lower for skip_word in ['total', 'percentage', 'percent', 'sum', outcome_col.lower()]):
                continue
            
            # Skip purely numeric rows (likely table artifacts)
            try:
                float(outcome_text)
                continue  # It's numeric, skip it
            except ValueError:
                pass  # It's text, keep it
            
            # Validate that at least some Bloom columns have numeric data for this row
            has_numeric_data = False
            outcome_id = len(outcomes)
            
            for bloom, col in bloom_columns.items():
                try:
                    val = row[col] if col in row.index else ""
                    if pd.isna(val) or val == "":
                        count = 0
                    else:
                        count = int(float(str(val).strip()))
                    tos_matrix[bloom][outcome_id] = count
                    total_items += count
                    if count > 0:
                        has_numeric_data = True
                except (ValueError, TypeError):
                    tos_matrix[bloom][outcome_id] = 0
            
            # Only add this outcome if it has at least some numeric data
            if has_numeric_data or idx == 0:  # Include first valid row even if empty counts
                outcomes.append({
                    "id": outcome_id,
                    "text": outcome_text,
                    "hours": 1.0
                })
        
        if not outcomes:
            raise TOSParsingError(
                f"No valid learning outcomes found in Excel sheet. "
                f"Checked {len(df)} rows after header for textual outcomes with Bloom data."
            )
        
        # STEP 9: Compute bloom distribution percentages
        bloom_distribution = {}
        for bloom in BLOOM_LEVELS:
            total_for_bloom = sum(tos_matrix[bloom].values())
            bloom_distribution[bloom] = total_for_bloom if total_items == 0 else int(
                (total_for_bloom / total_items) * 100
            )
        
        return {
            "learning_outcomes": outcomes,
            "bloom_distribution": bloom_distribution,
            "tos_matrix": tos_matrix,
            "total_items": total_items,
            "metadata": {
                "parsing_method": "xlsx_pandas_extraction"
            }
        }
        return {
            "learning_outcomes": outcomes,
            "bloom_distribution": bloom_distribution,
            "tos_matrix": tos_matrix,
            "total_items": total_items,
            "metadata": {
                "parsing_method": "xlsx_sheet_extraction"
            }
        }

    def _validate_and_normalize_tos(
        self,
        tos_data: Dict[str, Any],
        file_name: str
    ) -> Dict[str, Any]:
        """
        Validate TOS structure and normalize it to expected format.

        Ensures all required fields exist and are properly formatted.
        """
        # Validate required top-level fields
        required_fields = ["learning_outcomes", "bloom_distribution", "tos_matrix"]
        missing = [f for f in required_fields if f not in tos_data]
        if missing:
            raise TOSParsingError(f"Missing required fields: {missing}")

        # Validate learning outcomes
        outcomes = tos_data.get("learning_outcomes", [])
        if not outcomes:
            # Support in-app generated TOS structure
            outcomes = tos_data.get("outcomes", [])
        if not isinstance(outcomes, list) or not outcomes:
            raise TOSParsingError("learning_outcomes must be a non-empty list")

        for i, outcome in enumerate(outcomes):
            if not isinstance(outcome, dict):
                raise TOSParsingError(f"Outcome {i} is not a dict")
            if "text" not in outcome and "description" not in outcome:
                raise TOSParsingError(
                    f"Outcome {i} must have 'text' or 'description' field"
                )
            # Normalize: use 'text' as canonical field
            if "description" in outcome and "text" not in outcome:
                outcome["text"] = outcome.pop("description")
            # Ensure ID exists
            if "id" not in outcome:
                outcome["id"] = i

        # Validate Bloom distribution
        bloom_dist = tos_data.get("bloom_distribution", {})
        for bloom in BLOOM_LEVELS:
            if bloom not in bloom_dist:
                raise TOSParsingError(f"Missing Bloom level: {bloom}")
            val = bloom_dist[bloom]
            if not isinstance(val, (int, float)):
                raise TOSParsingError(f"Bloom '{bloom}' value must be numeric")

        # Validate TOS matrix structure
        tos_matrix = tos_data.get("tos_matrix", {})
        for bloom in BLOOM_LEVELS:
            if bloom not in tos_matrix:
                tos_matrix[bloom] = {}  # Add missing empty dicts
            matrix_row = tos_matrix[bloom]
            if not isinstance(matrix_row, dict):
                raise TOSParsingError(
                    f"TOS matrix entry for '{bloom}' must be a dict"
                )
            # Normalize matrix keys to integers for consistency
            normalized_row = {}
            for key, value in matrix_row.items():
                # Convert string keys to int if possible
                try:
                    int_key = int(key) if isinstance(key, str) else key
                except (ValueError, TypeError):
                    int_key = key
                normalized_row[int_key] = value
            tos_matrix[bloom] = normalized_row
            
            # Ensure all outcome IDs are in the row
            for outcome in outcomes:
                oid = outcome["id"]
                if oid not in tos_matrix[bloom]:
                    tos_matrix[bloom][oid] = 0

        # Compute total items
        total_items = sum(
            sum(tos_matrix[bloom].values())
            for bloom in BLOOM_LEVELS
        )
        tos_data["total_items"] = total_items

        # Add/update metadata
        if "metadata" not in tos_data:
            tos_data["metadata"] = {}
        tos_data["metadata"].update({
            "file_name": file_name,
            "parsed_at": datetime.now().isoformat(),
            "total_outcomes": len(outcomes)
        })

        return tos_data


def parse_tos_file(
    file_content: bytes,
    file_name: str,
    file_type: str = None
) -> Tuple[bool, Dict[str, Any]]:
    """
    Convenience function to parse a TOS file.

    Returns:
        Tuple of (success: bool, data: Dict)
        If successful: (True, normalized_tos_data)
        If failed: (False, {"error": error_message})
    """
    try:
        parser = TOSFileParser()
        result = parser.parse(file_content, file_name, file_type)
        return True, result
    except Exception as e:
        logger.error(f"TOS parsing error: {str(e)}")
        return False, {"error": str(e)}


def validate_tos_for_tqs_generation(tos_data: Dict[str, Any]) -> Tuple[bool, str]:
    """
    Validate that a parsed TOS is suitable for TQS generation.

    Returns:
        Tuple of (is_valid: bool, message: str)
    """
    try:
        # Check required fields
        if not tos_data.get("learning_outcomes"):
            return False, "No learning outcomes found in TOS"

        if not tos_data.get("tos_matrix"):
            return False, "No TOS matrix found"

        total_items = tos_data.get("total_items", 0)
        if total_items == 0:
            return False, "TOS contains no items (total_items = 0)"

        # Check Bloom distribution
        bloom_dist = tos_data.get("bloom_distribution", {})
        if not bloom_dist or all(v == 0 for v in bloom_dist.values()):
            return False, "Bloom distribution is empty"

        return True, "TOS is valid for TQS generation"

    except Exception as e:
        return False, f"Validation error: {str(e)}"


def convert_tos_to_assigned_slots(
    tos_data: Dict[str, Any],
    question_type: str,
    points_per_item: float = 1.0
) -> Tuple[bool, Any]:
    """
    Convert parsed TOS to assigned slots format for TQS generation.

    This bridges the file-uploaded TOS with the internal TQS generation pipeline.

    Args:
        tos_data: Parsed TOS structure
        question_type: Type of questions to generate ("MCQ", "Essay", etc.)
        points_per_item: Points per question (default 1.0)

    Returns:
        Tuple of (success: bool, assigned_slots_list or error_message)
    """
    try:
        outcomes = tos_data.get("learning_outcomes", [])
        if not outcomes:
            outcomes = tos_data.get("outcomes", [])
        tos_matrix = tos_data.get("tos_matrix", {})

        if not outcomes:
            return False, "No learning outcomes found in TOS"
        if not tos_matrix:
            return False, "No TOS matrix found"

        total_slots = 0
        for bloom_key, bloom_row in tos_matrix.items():
            if isinstance(bloom_row, dict):
                total_slots += sum(int(v) for v in bloom_row.values() if isinstance(v, (int, float)))
        if total_slots == 0:
            return False, "TOS contains no items (all slots are zero)"

        assigned_slots = []

        for bloom in BLOOM_LEVELS:
            bloom_row = tos_matrix.get(bloom, {})
            if not bloom_row:
                for key in tos_matrix.keys():
                    if str(key).lower() == bloom.lower():
                        bloom_row = tos_matrix.get(key, {})
                        break
            for outcome in outcomes:
                outcome_id = outcome.get("id", outcome.get("_id"))
                outcome_text = outcome.get("text", outcome.get("description"))
                # Get count, handling both int and string keys in bloom_row
                count = bloom_row.get(outcome_id, 0)
                # If not found and outcome_id is int, try string version
                if count == 0 and isinstance(outcome_id, int):
                    count = bloom_row.get(str(outcome_id), 0)
                # If not found and outcome_id is str, try int version
                elif count == 0 and isinstance(outcome_id, str):
                    try:
                        count = bloom_row.get(int(outcome_id), 0)
                    except (ValueError, TypeError):
                        pass
                # If still not found, try outcome_* style keys
                if count == 0 and outcome_id is not None:
                    count = bloom_row.get(f"outcome_{outcome_id}", 0)

                # Create a slot for each item
                for _ in range(count):
                    slot = {
                        "outcome_id": outcome_id,
                        "outcome_text": outcome_text,
                        "bloom_level": bloom,
                        "question_type": question_type,
                        "points": points_per_item
                    }
                    assigned_slots.append(slot)

        if not assigned_slots:
            return False, "No slots generated from TOS"

        return True, assigned_slots

    except Exception as e:
        logger.error(f"Error converting TOS to slots: {str(e)}")
        return False, f"Conversion error: {str(e)}"
