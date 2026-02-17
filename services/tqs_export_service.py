"""
TQS Export Service - Generate DOCX, PDF, and CSV files from test questions
"""

import io
import csv
import copy
import random
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY


class TQSExportService:
    """Service for exporting test questions to various formats."""
    
    def __init__(self):
        self.default_course_name = "Course Name"
        self.default_exam_title = "Test Question Sheet"
        self.default_instructions = [
            "Read all questions carefully before answering.",
            "Answer all questions.",
            "Write your answers clearly and legibly.",
            "Show all work for problem-solving questions."
        ]
    
    # ======================================================
    # EXPORT TO DOCX
    # ======================================================
    
    def export_to_docx(
        self, 
        questions: List[Dict[str, Any]], 
        course_name: str = None, 
        exam_title: str = None,
        instructions: List[str] = None,
        exam_term: str = "Midterm",
        instructor_name: str = "",
        shuffle_choices: bool = False,
        generate_versions: bool = False,
        num_versions: int = 2,
        shuffle_question_order: bool = True
    ) -> io.BytesIO:
        """
        Export questions to DOCX format.
        
        Args:
            questions: List of question dictionaries
            course_name: Name of the course
            exam_title: Title of the exam
            instructions: List of instruction strings
            exam_term: Exam term (Midterm, Final, etc.)
            instructor_name: Instructor name
            shuffle_choices: Whether to shuffle MCQ choices
            generate_versions: Whether to generate multiple versions (A, B, etc.)
            num_versions: Number of versions to generate (default: 2)
            shuffle_question_order: Whether to shuffle question order in versions
            
        Returns:
            BytesIO object containing DOCX file
        """
        # Generate versions if requested
        if generate_versions:
            return self._export_docx_with_versions(
                questions=questions,
                course_name=course_name,
                exam_title=exam_title,
                instructions=instructions,
                exam_term=exam_term,
                instructor_name=instructor_name,
                num_versions=num_versions,
                shuffle_question_order=shuffle_question_order
            )
        
        # Apply choice shuffling if requested (single version)
        if shuffle_choices:
            questions = self.shuffle_questions_choices(questions)
        
        # Generate single version document
        return self._generate_single_docx(
            questions=questions,
            course_name=course_name,
            exam_title=exam_title,
            instructions=instructions,
            exam_term=exam_term,
            instructor_name=instructor_name
        )
    
    def _generate_single_docx(
        self,
        questions: List[Dict[str, Any]], 
        course_name: str = None, 
        exam_title: str = None,
        instructions: List[str] = None,
        exam_term: str = "Midterm",
        instructor_name: str = "",
        version_label: str = None
    ) -> io.BytesIO:
        """Generate a single DOCX document."""
        doc = Document()
        
        # Set default styles
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Add Header with version label if provided
        title_text = exam_title or self.default_exam_title
        if version_label:
            title_text = f"{title_text} - {version_label}"
        
        header = doc.add_heading(title_text, level=0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add Course Info
        course_para = doc.add_paragraph()
        course_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        course_run = course_para.add_run(course_name or self.default_course_name)
        course_run.bold = True
        course_run.font.size = Pt(12)
        
        # Add Exam Term
        if exam_term:
            term_para = doc.add_paragraph()
            term_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            term_run = term_para.add_run(exam_term)
            term_run.font.size = Pt(11)
        
        # Add Instructor
        if instructor_name:
            instructor_para = doc.add_paragraph()
            instructor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            instructor_run = instructor_para.add_run(f"Instructor: {instructor_name}")
            instructor_run.font.size = Pt(10)
        
        # Add Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(10)
        
        doc.add_paragraph()  # Spacing
        
        # Add Instructions
        doc.add_heading('Instructions:', level=2)
        instruction_list = instructions or self.default_instructions
        for instruction in instruction_list:
            doc.add_paragraph(instruction, style='List Bullet')
        
        doc.add_paragraph()  # Spacing
        
        # Add Total Points
        total_points = sum(q.get('points', 1) for q in questions)
        points_para = doc.add_paragraph()
        points_run = points_para.add_run(f"Total Points: {total_points}")
        points_run.bold = True
        points_run.font.size = Pt(11)
        
        doc.add_paragraph('_' * 80)  # Horizontal line
        doc.add_paragraph()
        
        # Add Questions
        for q in questions:
            self._add_question_to_docx(doc, q)
        
        # Add Page Break before Answer Key
        doc.add_page_break()
        
        # Add Answer Key
        self._add_answer_key_to_docx(doc, questions)
        
        # Save to BytesIO
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        return docx_file
    
    def _export_docx_with_versions(
        self,
        questions: List[Dict[str, Any]], 
        course_name: str = None, 
        exam_title: str = None,
        instructions: List[str] = None,
        exam_term: str = "Midterm",
        instructor_name: str = "",
        num_versions: int = 2,
        shuffle_question_order: bool = True
    ) -> io.BytesIO:
        """
        Generate multiple exam versions in a single DOCX document.
        
        Each version has its own questions and answer key section.
        """
        # Generate versions
        versions = self.generate_exam_versions(
            questions=questions,
            num_versions=num_versions,
            shuffle_question_order=shuffle_question_order,
            shuffle_choices=True
        )
        
        doc = Document()
        
        # Set default styles
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Generate each version
        for version_idx, (version_label, version_questions) in enumerate(versions):
            # Add version separator if not first version
            if version_idx > 0:
                doc.add_page_break()
            
            # Add Header
            title_text = f"{exam_title or self.default_exam_title} - {version_label}"
            header = doc.add_heading(title_text, level=0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add Course Info
            course_para = doc.add_paragraph()
            course_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            course_run = course_para.add_run(course_name or self.default_course_name)
            course_run.bold = True
            course_run.font.size = Pt(12)
            
            # Add Exam Term
            if exam_term:
                term_para = doc.add_paragraph()
                term_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                term_run = term_para.add_run(exam_term)
                term_run.font.size = Pt(11)
            
            # Add Instructor
            if instructor_name:
                instructor_para = doc.add_paragraph()
                instructor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                instructor_run = instructor_para.add_run(f"Instructor: {instructor_name}")
                instructor_run.font.size = Pt(10)
            
            # Add Date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
            date_run.font.size = Pt(10)
            
            doc.add_paragraph()  # Spacing
            
            # Add Instructions
            doc.add_heading('Instructions:', level=2)
            instruction_list = instructions or self.default_instructions
            for instruction in instruction_list:
                doc.add_paragraph(instruction, style='List Bullet')
            
            doc.add_paragraph()  # Spacing
            
            # Add Total Points
            total_points = sum(q.get('points', 1) for q in version_questions)
            points_para = doc.add_paragraph()
            points_run = points_para.add_run(f"Total Points: {total_points}")
            points_run.bold = True
            points_run.font.size = Pt(11)
            
            doc.add_paragraph('_' * 80)  # Horizontal line
            doc.add_paragraph()
            
            # Add Questions for this version
            for q in version_questions:
                self._add_question_to_docx(doc, q)
            
            # Add Page Break before Answer Key
            doc.add_page_break()
            
            # Add Answer Key for this version
            answer_key_title = f"ANSWER KEY - {version_label}"
            header = doc.add_heading(answer_key_title, level=1)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph('_' * 80)
            doc.add_paragraph()
            
            # Add answers for this version
            for q in version_questions:
                q_num = q.get('question_number', 1)
                q_type = q.get('type', q.get('question_type', 'MCQ'))
                bloom = q.get('bloom_level', q.get('bloom', 'Remember'))
                points = q.get('points', 1)
                
                answer_para = doc.add_paragraph()
                
                # Question number with Bloom level
                answer_para.add_run(f"Question {q_num}: ").bold = True
                
                # Answer
                if q_type == 'MCQ':
                    correct_answer = q.get('correct_answer', 'A')
                    answer_para.add_run(f"{correct_answer}")
                    
                    # Show full answer text
                    choices = q.get('choices', [])
                    choice_labels = ['A', 'B', 'C', 'D']
                    if correct_answer in choice_labels:
                        idx = choice_labels.index(correct_answer)
                        if idx < len(choices):
                            answer_para.add_run(f" - {choices[idx]}")
                
                elif q_type == 'Short Answer':
                    answer_key = q.get('answer_key', q.get('sample_answer', 'See instructor rubric'))
                    answer_para.add_run(answer_key)
                
                else:  # Essay, Problem Solving, Drawing
                    sample_answer = q.get('sample_answer', q.get('answer_key', 'See grading rubric'))
                    answer_para.add_run(sample_answer[:100] + "..." if len(sample_answer) > 100 else sample_answer)
                
                # Add metadata with bloom level
                answer_para.add_run(f"  [Bloom: {bloom} | {points} pts]")
                
                doc.add_paragraph()  # Spacing
        
        # Save to BytesIO
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        return docx_file
    
    def _add_question_to_docx(self, doc: Document, question: Dict[str, Any]):
        """Add a single question to the DOCX document."""
        q_num = question.get('question_number', 1)
        q_text = question.get('question_text', '')
        q_type = question.get('type', question.get('question_type', 'MCQ'))
        q_points = question.get('points', 1)
        
        # Question header
        q_header = doc.add_paragraph()
        q_header_run = q_header.add_run(f"Question {q_num}. ({q_points} {'point' if q_points == 1 else 'points'})")
        q_header_run.bold = True
        q_header_run.font.size = Pt(11)
        
        # Question text
        q_para = doc.add_paragraph(q_text)
        q_para.paragraph_format.left_indent = Inches(0.25)
        
        # Add type-specific content
        if q_type == 'MCQ':
            choices = question.get('choices', [])
            choice_labels = ['A', 'B', 'C', 'D']
            for i, choice in enumerate(choices[:4]):
                choice_para = doc.add_paragraph()
                choice_para.paragraph_format.left_indent = Inches(0.5)
                choice_text = f"{choice_labels[i]}. {choice}"
                choice_para.add_run(choice_text)
        
        elif q_type == 'Short Answer':
            answer_para = doc.add_paragraph()
            answer_para.paragraph_format.left_indent = Inches(0.5)
            answer_para.add_run("Answer: " + "_" * 60)
        
        elif q_type in ['Essay', 'Problem Solving']:
            answer_para = doc.add_paragraph()
            answer_para.paragraph_format.left_indent = Inches(0.5)
            answer_para.add_run("Answer space:\n\n\n\n")
        
        doc.add_paragraph()  # Spacing between questions
    
    def _add_answer_key_to_docx(self, doc: Document, questions: List[Dict[str, Any]]):
        """Add answer key section to DOCX document."""
        # Answer Key Header
        header = doc.add_heading('ANSWER KEY', level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
        
        # Add answers
        for q in questions:
            q_num = q.get('question_number', 1)
            q_type = q.get('type', q.get('question_type', 'MCQ'))
            bloom = q.get('bloom_level', q.get('bloom', 'Remember'))
            points = q.get('points', 1)
            
            answer_para = doc.add_paragraph()
            
            # Question number
            answer_para.add_run(f"Question {q_num}: ").bold = True
            
            # Answer
            if q_type == 'MCQ':
                correct_answer = q.get('correct_answer', 'A')
                answer_para.add_run(correct_answer)
                
                # Show full answer text
                choices = q.get('choices', [])
                choice_labels = ['A', 'B', 'C', 'D']
                if correct_answer in choice_labels:
                    idx = choice_labels.index(correct_answer)
                    if idx < len(choices):
                        answer_para.add_run(f" - {choices[idx]}")
            
            elif q_type == 'Short Answer':
                answer_key = q.get('answer_key', q.get('sample_answer', 'See instructor rubric'))
                answer_para.add_run(answer_key)
            
            else:  # Essay, Problem Solving, Drawing
                sample_answer = q.get('sample_answer', q.get('answer_key', 'See grading rubric'))
                answer_para.add_run(sample_answer[:100] + "..." if len(sample_answer) > 100 else sample_answer)
            
            # Add metadata
            meta_para = doc.add_paragraph()
            meta_para.paragraph_format.left_indent = Inches(0.25)
            meta_run = meta_para.add_run(f"  Bloom Level: {bloom} | Points: {points}")
            meta_run.font.size = Pt(9)
            meta_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()  # Spacing
    
    # ======================================================
    # EXPORT TO PDF
    # ======================================================
    
    def export_to_pdf(
        self, 
        questions: List[Dict[str, Any]], 
        course_name: str = None, 
        exam_title: str = None,
        instructions: List[str] = None,
        exam_term: str = "Midterm",
        instructor_name: str = "",
        shuffle_choices: bool = False,
        generate_versions: bool = False,
        num_versions: int = 2,
        shuffle_question_order: bool = True
    ) -> io.BytesIO:
        """
        Export questions to PDF format.
        
        Args:
            questions: List of question dictionaries
            course_name: Name of the course
            exam_title: Title of the exam
            instructions: List of instruction strings
            exam_term: Exam term
            instructor_name: Instructor name
            shuffle_choices: Whether to shuffle MCQ choices
            generate_versions: Whether to generate multiple versions (A, B, etc.)
            num_versions: Number of versions to generate (default: 2)
            shuffle_question_order: Whether to shuffle question order in versions
            
        Returns:
            BytesIO object containing PDF file
        """
        # Generate versions if requested
        if generate_versions:
            return self._export_pdf_with_versions(
                questions=questions,
                course_name=course_name,
                exam_title=exam_title,
                instructions=instructions,
                exam_term=exam_term,
                instructor_name=instructor_name,
                num_versions=num_versions,
                shuffle_question_order=shuffle_question_order
            )
        
        # Apply choice shuffling if requested (single version)
        if shuffle_choices:
            questions = self.shuffle_questions_choices(questions)
        
        # Generate single version PDF
        return self._generate_single_pdf(
            questions=questions,
            course_name=course_name,
            exam_title=exam_title,
            instructions=instructions,
            exam_term=exam_term,
            instructor_name=instructor_name
        )
    
    def _generate_single_pdf(
        self,
        questions: List[Dict[str, Any]], 
        course_name: str = None, 
        exam_title: str = None,
        instructions: List[str] = None,
        exam_term: str = "Midterm",
        instructor_name: str = "",
        version_label: str = None
    ) -> io.BytesIO:
        """Generate a single PDF document."""
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            pdf_buffer, 
            pagesize=letter,
            leftMargin=0.75*inch,
            rightMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        # Container for PDF elements
        elements = []
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#333333'),
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        info_style = ParagraphStyle(
            'InfoStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#666666'),
            spaceAfter=4,
            alignment=TA_CENTER,
            fontName='Helvetica'
        )
        
        question_style = ParagraphStyle(
            'QuestionStyle',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.black,
            spaceAfter=8,
            alignment=TA_JUSTIFY,
            fontName='Helvetica'
        )
        
        # Add Header with version label if provided
        title_text = exam_title or self.default_exam_title
        if version_label:
            title_text = f"{title_text} - {version_label}"
        
        elements.append(Paragraph(title_text, title_style))
        elements.append(Paragraph(course_name or self.default_course_name, subtitle_style))
        
        if exam_term:
            elements.append(Paragraph(exam_term, info_style))
        
        if instructor_name:
            elements.append(Paragraph(f"Instructor: {instructor_name}", info_style))
        
        elements.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", info_style))
        elements.append(Spacer(1, 0.3*inch))
        
        # Add Instructions
        elements.append(Paragraph("<b>Instructions:</b>", styles['Heading2']))
        instruction_list = instructions or self.default_instructions
        for instruction in instruction_list:
            elements.append(Paragraph(f"• {instruction}", question_style))
        
        elements.append(Spacer(1, 0.2*inch))
        
        # Add Total Points
        total_points = sum(q.get('points', 1) for q in questions)
        elements.append(Paragraph(f"<b>Total Points: {total_points}</b>", question_style))
        elements.append(Spacer(1, 0.1*inch))
        
        # Horizontal line
        elements.append(Paragraph("_" * 100, styles['Normal']))
        elements.append(Spacer(1, 0.2*inch))
        
        # Add Questions
        for q in questions:
            self._add_question_to_pdf(elements, q, question_style, styles)
        
        # Page Break before Answer Key
        elements.append(PageBreak())
        
        # Add Answer Key
        self._add_answer_key_to_pdf(elements, questions, title_style, question_style, styles)
        
        # Build PDF
        doc.build(elements)
        pdf_buffer.seek(0)
        
        return pdf_buffer
    
    def _export_pdf_with_versions(
        self,
        questions: List[Dict[str, Any]], 
        course_name: str = None, 
        exam_title: str = None,
        instructions: List[str] = None,
        exam_term: str = "Midterm",
        instructor_name: str = "",
        num_versions: int = 2,
        shuffle_question_order: bool = True
    ) -> io.BytesIO:
        """
        Generate multiple exam versions in a single PDF document.
        
        Each version has its own questions and answer key section.
        """
        # Generate versions
        versions = self.generate_exam_versions(
            questions=questions,
            num_versions=num_versions,
            shuffle_question_order=shuffle_question_order,
            shuffle_choices=True
        )
        
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            pdf_buffer, 
            pagesize=letter,
            leftMargin=0.75*inch,
            rightMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        elements = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#333333'),
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        info_style = ParagraphStyle(
            'InfoStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#666666'),
            spaceAfter=4,
            alignment=TA_CENTER,
            fontName='Helvetica'
        )
        
        question_style = ParagraphStyle(
            'QuestionStyle',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.black,
            spaceAfter=8,
            alignment=TA_JUSTIFY,
            fontName='Helvetica'
        )
        
        # Generate each version
        for version_idx, (version_label, version_questions) in enumerate(versions):
            # Add page break if not first version
            if version_idx > 0:
                elements.append(PageBreak())
            
            # Add Header
            title_text = f"{exam_title or self.default_exam_title} - {version_label}"
            elements.append(Paragraph(title_text, title_style))
            elements.append(Paragraph(course_name or self.default_course_name, subtitle_style))
            
            if exam_term:
                elements.append(Paragraph(exam_term, info_style))
            
            if instructor_name:
                elements.append(Paragraph(f"Instructor: {instructor_name}", info_style))
            
            elements.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", info_style))
            elements.append(Spacer(1, 0.3*inch))
            
            # Add Instructions
            elements.append(Paragraph("<b>Instructions:</b>", styles['Heading2']))
            instruction_list = instructions or self.default_instructions
            for instruction in instruction_list:
                elements.append(Paragraph(f"• {instruction}", question_style))
            
            elements.append(Spacer(1, 0.2*inch))
            
            # Add Total Points
            total_points = sum(q.get('points', 1) for q in version_questions)
            elements.append(Paragraph(f"<b>Total Points: {total_points}</b>", question_style))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("_" * 100, styles['Normal']))
            elements.append(Spacer(1, 0.2*inch))
            
            # Add Questions for this version
            for q in version_questions:
                self._add_question_to_pdf(elements, q, question_style, styles)
            
            # Page Break before Answer Key
            elements.append(PageBreak())
            
            # Add Answer Key for this version
            answer_key_title = f"ANSWER KEY - {version_label}"
            elements.append(Paragraph(answer_key_title, title_style))
            elements.append(Paragraph("_" * 100, styles['Normal']))
            elements.append(Spacer(1, 0.2*inch))
            
            # Add answers for this version
            for q in version_questions:
                q_num = q.get('question_number', 1)
                q_type = q.get('type', q.get('question_type', 'MCQ'))
                bloom = q.get('bloom_level', q.get('bloom', 'Remember'))
                points = q.get('points', 1)
                
                # Build answer text
                answer_text = ""
                if q_type == 'MCQ':
                    correct_answer = q.get('correct_answer', 'A')
                    choices = q.get('choices', [])
                    choice_labels = ['A', 'B', 'C', 'D']
                    
                    if correct_answer in choice_labels:
                        idx = choice_labels.index(correct_answer)
                        if idx < len(choices):
                            answer_text = f"<b>{correct_answer}</b> - {choices[idx]}"
                        else:
                            answer_text = f"<b>{correct_answer}</b>"
                    else:
                        answer_text = f"<b>{correct_answer}</b>"
                
                elif q_type == 'Short Answer':
                    answer_key = q.get('answer_key', q.get('sample_answer', 'See instructor rubric'))
                    answer_text = answer_key[:100] + "..." if len(answer_key) > 100 else answer_key
                
                else:
                    sample_answer = q.get('sample_answer', q.get('answer_key', 'See grading rubric'))
                    answer_text = sample_answer[:100] + "..." if len(sample_answer) > 100 else sample_answer
                
                # Add to PDF with Bloom level
                elements.append(Paragraph(
                    f"<b>Question {q_num}:</b> {answer_text} <font size=9 color='#666666'>[Bloom: {bloom} | {points} pts]</font>", 
                    question_style
                ))
                elements.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(elements)
        pdf_buffer.seek(0)
        
        return pdf_buffer
    
    def _add_question_to_pdf(self, elements: list, question: Dict[str, Any], question_style, styles):
        """Add a single question to the PDF document."""
        q_num = question.get('question_number', 1)
        q_text = question.get('question_text', '')
        q_type = question.get('type', question.get('question_type', 'MCQ'))
        q_points = question.get('points', 1)
        
        # Question header
        q_header = f"<b>Question {q_num}. ({q_points} {'point' if q_points == 1 else 'points'})</b>"
        elements.append(Paragraph(q_header, question_style))
        
        # Question text
        elements.append(Paragraph(q_text, question_style))
        elements.append(Spacer(1, 0.1*inch))
        
        # Add type-specific content
        if q_type == 'MCQ':
            choices = question.get('choices', [])
            choice_labels = ['A', 'B', 'C', 'D']
            for i, choice in enumerate(choices[:4]):
                choice_text = f"&nbsp;&nbsp;&nbsp;&nbsp;{choice_labels[i]}. {choice}"
                elements.append(Paragraph(choice_text, question_style))
        
        elif q_type == 'Short Answer':
            elements.append(Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;Answer: _______________________________________________", question_style))
        
        elif q_type in ['Essay', 'Problem Solving']:
            elements.append(Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;Answer space:", question_style))
            elements.append(Spacer(1, 0.5*inch))
        
        elements.append(Spacer(1, 0.2*inch))
    
    def _add_answer_key_to_pdf(self, elements: list, questions: List[Dict[str, Any]], title_style, question_style, styles):
        """Add answer key section to PDF document."""
        # Answer Key Header
        elements.append(Paragraph("ANSWER KEY", title_style))
        elements.append(Paragraph("_" * 100, styles['Normal']))
        elements.append(Spacer(1, 0.2*inch))
        
        # Create answer key data
        for q in questions:
            q_num = q.get('question_number', 1)
            q_type = q.get('type', q.get('question_type', 'MCQ'))
            bloom = q.get('bloom_level', q.get('bloom', 'Remember'))
            points = q.get('points', 1)
            
            # Build answer text
            answer_text = ""
            if q_type == 'MCQ':
                correct_answer = q.get('correct_answer', 'A')
                choices = q.get('choices', [])
                choice_labels = ['A', 'B', 'C', 'D']
                
                if correct_answer in choice_labels:
                    idx = choice_labels.index(correct_answer)
                    if idx < len(choices):
                        answer_text = f"<b>{correct_answer}</b> - {choices[idx]}"
                    else:
                        answer_text = f"<b>{correct_answer}</b>"
                else:
                    answer_text = f"<b>{correct_answer}</b>"
            
            elif q_type == 'Short Answer':
                answer_key = q.get('answer_key', q.get('sample_answer', 'See instructor rubric'))
                answer_text = answer_key[:100] + "..." if len(answer_key) > 100 else answer_key
            
            else:  # Essay, Problem Solving, Drawing
                sample_answer = q.get('sample_answer', q.get('answer_key', 'See grading rubric'))
                answer_text = sample_answer[:100] + "..." if len(sample_answer) > 100 else sample_answer
            
            # Add to PDF
            elements.append(Paragraph(f"<b>Question {q_num}:</b> {answer_text}", question_style))
            
            # Add metadata
            meta_text = f"<font size=9 color='#666666'>&nbsp;&nbsp;&nbsp;&nbsp;Bloom Level: {bloom} | Points: {points}</font>"
            elements.append(Paragraph(meta_text, question_style))
            elements.append(Spacer(1, 0.1*inch))
    
    # ======================================================
    # EXPORT TO CSV
    # ======================================================
    
    def export_to_csv(self, questions: List[Dict[str, Any]]) -> io.StringIO:
        """
        Export questions to CSV format.
        
        Format: Question, Option A, Option B, Option C, Option D, Correct Answer, Bloom Level, Points
        
        Args:
            questions: List of question dictionaries
            
        Returns:
            StringIO object containing CSV data
        """
        csv_buffer = io.StringIO()
        
        # Define CSV headers
        fieldnames = [
            'Question Number',
            'Question Text',
            'Question Type',
            'Option A',
            'Option B',
            'Option C',
            'Option D',
            'Correct Answer',
            'Answer Key/Sample Answer',
            'Bloom Level',
            'Points',
            'Learning Outcome'
        ]
        
        writer = csv.DictWriter(csv_buffer, fieldnames=fieldnames)
        writer.writeheader()
        
        # Write questions
        for q in questions:
            q_type = q.get('type', q.get('question_type', 'MCQ'))
            choices = q.get('choices', [])
            
            # Build row data
            row = {
                'Question Number': q.get('question_number', ''),
                'Question Text': q.get('question_text', ''),
                'Question Type': q_type,
                'Option A': choices[0] if len(choices) > 0 else '',
                'Option B': choices[1] if len(choices) > 1 else '',
                'Option C': choices[2] if len(choices) > 2 else '',
                'Option D': choices[3] if len(choices) > 3 else '',
                'Correct Answer': q.get('correct_answer', '') if q_type == 'MCQ' else 'N/A',
                'Answer Key/Sample Answer': q.get('answer_key', q.get('sample_answer', '')),
                'Bloom Level': q.get('bloom_level', q.get('bloom', '')),
                'Points': q.get('points', 1),
                'Learning Outcome': q.get('outcome_text', q.get('learning_outcome', ''))
            }
            
            writer.writerow(row)
        
        csv_buffer.seek(0)
        return csv_buffer
    
    # ======================================================
    # UTILITY METHODS
    # ======================================================
    
    def get_exam_metadata(self, questions: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Get metadata about the exam."""
        total_points = sum(q.get('points', 1) for q in questions)
        
        question_types = {}
        bloom_levels = {}
        
        for q in questions:
            q_type = q.get('type', q.get('question_type', 'Unknown'))
            bloom = q.get('bloom_level', q.get('bloom', 'Unknown'))
            
            question_types[q_type] = question_types.get(q_type, 0) + 1
            bloom_levels[bloom] = bloom_levels.get(bloom, 0) + 1
        
        return {
            'total_questions': len(questions),
            'total_points': total_points,
            'question_types': question_types,
            'bloom_levels': bloom_levels
        }
    
    # ======================================================
    # CHOICE SHUFFLING & VERSION GENERATION
    # ======================================================
    
    def _shuffle_mcq_choices(self, question: Dict[str, Any], seed: int = None) -> Dict[str, Any]:
        """
        Shuffle MCQ choices and update correct_answer accordingly.
        
        Args:
            question: Question dictionary (will be copied, not modified)
            seed: Random seed for reproducible shuffling
            
        Returns:
            New question dictionary with shuffled choices
        """
        # Create a deep copy to avoid modifying original
        q_copy = copy.deepcopy(question)
        
        q_type = q_copy.get('type', q_copy.get('question_type', ''))
        
        # Only shuffle MCQ questions
        if q_type != 'MCQ':
            return q_copy
        
        choices = q_copy.get('choices', [])
        if len(choices) < 2:
            return q_copy
        
        correct_answer = q_copy.get('correct_answer', 'A')
        choice_labels = ['A', 'B', 'C', 'D']
        
        # Get the index of correct answer
        if correct_answer not in choice_labels:
            return q_copy
        
        correct_index = choice_labels.index(correct_answer)
        
        # Create list of (choice, is_correct) tuples
        choice_pairs = [(choices[i], i == correct_index) for i in range(len(choices))]
        
        # Shuffle with seed if provided
        if seed is not None:
            random.seed(seed)
        random.shuffle(choice_pairs)
        
        # Update choices and find new correct answer position
        new_choices = []
        new_correct_index = 0
        for i, (choice, is_correct) in enumerate(choice_pairs):
            new_choices.append(choice)
            if is_correct:
                new_correct_index = i
        
        q_copy['choices'] = new_choices
        q_copy['correct_answer'] = choice_labels[new_correct_index]
        
        return q_copy
    
    def shuffle_questions_choices(
        self, 
        questions: List[Dict[str, Any]], 
        seed: int = None
    ) -> List[Dict[str, Any]]:
        """
        Shuffle choices for all MCQ questions in the list.
        
        Args:
            questions: List of questions (will be copied, not modified)
            seed: Random seed for reproducible shuffling
            
        Returns:
            New list with shuffled MCQ choices
        """
        shuffled_questions = []
        
        for i, q in enumerate(questions):
            # Use different seed for each question if seed provided
            question_seed = None if seed is None else seed + i
            shuffled_q = self._shuffle_mcq_choices(q, question_seed)
            shuffled_questions.append(shuffled_q)
        
        return shuffled_questions
    
    def generate_exam_versions(
        self, 
        questions: List[Dict[str, Any]], 
        num_versions: int = 2,
        shuffle_question_order: bool = True,
        shuffle_choices: bool = True
    ) -> List[Tuple[str, List[Dict[str, Any]]]]:
        """
        Generate multiple versions of the exam with different orders and/or shuffled choices.
        
        Args:
            questions: Original list of questions (will not be modified)
            num_versions: Number of versions to generate (default: 2)
            shuffle_question_order: Whether to shuffle question order
            shuffle_choices: Whether to shuffle MCQ choices
            
        Returns:
            List of tuples: [(version_label, questions_list), ...]
            Example: [("Version A", questions_a), ("Version B", questions_b)]
        """
        versions = []
        version_labels = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
        
        for i in range(num_versions):
            version_label = version_labels[i] if i < len(version_labels) else f"Version {i+1}"
            
            # Create deep copy of questions
            version_questions = copy.deepcopy(questions)
            
            # Shuffle question order if enabled
            if shuffle_question_order:
                random.seed(1000 + i)  # Different seed for each version
                random.shuffle(version_questions)
                
                # Renumber questions
                for idx, q in enumerate(version_questions):
                    q['question_number'] = idx + 1
            
            # Shuffle MCQ choices if enabled
            if shuffle_choices:
                version_questions = self.shuffle_questions_choices(
                    version_questions, 
                    seed=2000 + i
                )
            
            versions.append((version_label, version_questions))
        
        return versions


# Singleton instance
tqs_export_service = TQSExportService()
